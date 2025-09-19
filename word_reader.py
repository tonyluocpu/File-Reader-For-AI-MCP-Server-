#!/usr/bin/env python3
"""
Word document reader module for AI File Bridge
Extracts text content from Microsoft Word documents (.docx files)
"""

import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


class WordReader:
    """Word document reader with formatting preservation"""
    
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def read_word_document(self, file_path: str, include_formatting: bool = False) -> Dict[str, Any]:
        """
        Read a Word document and extract its content
        
        Args:
            file_path: Path to .docx file
            include_formatting: Whether to include formatting information
            
        Returns:
            Dictionary with extracted Word document data
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not path.suffix.lower() == '.docx':
            raise ValueError(f"File must be a .docx file, got: {path.suffix}")
        
        try:
            doc = Document(path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraph_data = {"text": para.text.strip()}
                    
                    if include_formatting:
                        # Basic formatting info
                        paragraph_data["formatting"] = {
                            "style": para.style.name if para.style else "Normal",
                            "alignment": str(para.alignment) if para.alignment else None
                        }
                        
                        # Check for bold, italic in runs
                        runs_info = []
                        for run in para.runs:
                            if run.text.strip():
                                runs_info.append({
                                    "text": run.text.strip(),
                                    "bold": run.bold,
                                    "italic": run.italic,
                                    "underline": run.underline
                                })
                        if runs_info:
                            paragraph_data["runs"] = runs_info
                    
                    paragraphs.append(paragraph_data)
            
            # Extract tables if any
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "success": True,
                "file_path": str(path),
                "file_type": "Microsoft Word Document (.docx)",
                "paragraphs": paragraphs,
                "tables": tables,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "formatting_included": include_formatting
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_path": str(path)
            }


def main():
    """Test function for Word reader"""
    reader = WordReader()
    
    # Test with a sample file if available
    test_file = "sample.docx"
    if Path(test_file).exists():
        result = reader.read_word_document(test_file, include_formatting=True)
        print(json.dumps(result, indent=2, ensure_ascii=False, default=str))
    else:
        print(f"Test file {test_file} not found. Create a sample Word document to test.")


if __name__ == "__main__":
    import json
    main()
