#!/usr/bin/env python3
"""
MCP Server for reading Microsoft Word documents on macOS
Uses python-docx library to extract text content from .docx files
"""

import asyncio
import json
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


class WordReaderMCPServer:
    def __init__(self):
        self.server_name = "word-reader-mcp"
        self.version = "1.0.0"
    
    async def handle_request(self, request: Dict[str, Any]) -> Dict[str, Any]:
        """Handle incoming MCP requests"""
        method = request.get("method")
        params = request.get("params", {})
        
        if method == "initialize":
            return await self._handle_initialize(params)
        elif method == "tools/list":
            return await self._handle_tools_list()
        elif method == "tools/call":
            return await self._handle_tools_call(params)
        else:
            return {
                "jsonrpc": "2.0",
                "id": request.get("id"),
                "error": {
                    "code": -32601,
                    "message": f"Method not found: {method}"
                }
            }
    
    async def _handle_initialize(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Handle initialization request"""
        return {
            "jsonrpc": "2.0",
            "id": params.get("id"),
            "result": {
                "protocolVersion": "2024-11-05",
                "capabilities": {
                    "tools": {}
                },
                "serverInfo": {
                    "name": self.server_name,
                    "version": self.version
                }
            }
        }
    
    async def _handle_tools_list(self) -> Dict[str, Any]:
        """Return list of available tools"""
        return {
            "jsonrpc": "2.0",
            "result": {
                "tools": [
                    {
                        "name": "read_word_document",
                        "description": "Read and extract text content from a Microsoft Word document (.docx file)",
                        "inputSchema": {
                            "type": "object",
                            "properties": {
                                "file_path": {
                                    "type": "string",
                                    "description": "Path to the .docx file to read"
                                },
                                "include_formatting": {
                                    "type": "boolean",
                                    "description": "Whether to include basic formatting information",
                                    "default": False
                                }
                            },
                            "required": ["file_path"]
                        }
                    },
                    {
                        "name": "list_word_documents",
                        "description": "List all .docx files in a directory",
                        "inputSchema": {
                            "type": "object",
                            "properties": {
                                "directory_path": {
                                    "type": "string",
                                    "description": "Path to the directory to search for .docx files",
                                    "default": "."
                                }
                            }
                        }
                    }
                ]
            }
        }
    
    async def _handle_tools_call(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Handle tool call requests"""
        tool_name = params.get("name")
        arguments = params.get("arguments", {})
        
        try:
            if tool_name == "read_word_document":
                result = await self._read_word_document(arguments)
            elif tool_name == "list_word_documents":
                result = await self._list_word_documents(arguments)
            else:
                return {
                    "jsonrpc": "2.0",
                    "id": params.get("id"),
                    "error": {
                        "code": -32601,
                        "message": f"Tool not found: {tool_name}"
                    }
                }
            
            return {
                "jsonrpc": "2.0",
                "id": params.get("id"),
                "result": {
                    "content": [
                        {
                            "type": "text",
                            "text": json.dumps(result, indent=2, ensure_ascii=False)
                        }
                    ]
                }
            }
        except Exception as e:
            return {
                "jsonrpc": "2.0",
                "id": params.get("id"),
                "error": {
                    "code": -32603,
                    "message": f"Internal error: {str(e)}"
                }
            }
    
    async def _read_word_document(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Read a Word document and extract its content"""
        file_path = arguments.get("file_path")
        include_formatting = arguments.get("include_formatting", False)
        
        if not file_path:
            raise ValueError("file_path is required")
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not path.suffix.lower() == '.docx':
            raise ValueError(f"File must be a .docx file, got: {path.suffix}")
        
        try:
            doc = Document(path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraph_data = {"text": para.text.strip()}
                    
                    if include_formatting:
                        # Basic formatting info
                        paragraph_data["formatting"] = {
                            "style": para.style.name if para.style else "Normal",
                            "alignment": str(para.alignment) if para.alignment else None
                        }
                        
                        # Check for bold, italic in runs
                        runs_info = []
                        for run in para.runs:
                            if run.text.strip():
                                runs_info.append({
                                    "text": run.text.strip(),
                                    "bold": run.bold,
                                    "italic": run.italic,
                                    "underline": run.underline
                                })
                        if runs_info:
                            paragraph_data["runs"] = runs_info
                    
                    paragraphs.append(paragraph_data)
            
            # Extract tables if any
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "success": True,
                "file_path": str(path),
                "paragraphs": paragraphs,
                "tables": tables,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_path": str(path)
            }
    
    async def _list_word_documents(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """List all .docx files in a directory"""
        directory_path = arguments.get("directory_path", ".")
        path = Path(directory_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if not path.is_dir():
            raise ValueError(f"Path is not a directory: {directory_path}")
        
        docx_files = []
        for file_path in path.glob("*.docx"):
            docx_files.append({
                "name": file_path.name,
                "path": str(file_path),
                "size": file_path.stat().st_size,
                "modified": file_path.stat().st_mtime
            })
        
        return {
            "success": True,
            "directory": str(path),
            "docx_files": docx_files,
            "count": len(docx_files)
        }


async def main():
    """Main entry point for the MCP server"""
    server = WordReaderMCPServer()
    
    # Read from stdin and write to stdout
    while True:
        try:
            line = await asyncio.get_event_loop().run_in_executor(None, sys.stdin.readline)
            if not line:
                break
            
            request = json.loads(line.strip())
            response = await server.handle_request(request)
            print(json.dumps(response, ensure_ascii=False))
            sys.stdout.flush()
            
        except json.JSONDecodeError as e:
            error_response = {
                "jsonrpc": "2.0",
                "id": None,
                "error": {
                    "code": -32700,
                    "message": f"Parse error: {str(e)}"
                }
            }
            print(json.dumps(error_response))
            sys.stdout.flush()
        except Exception as e:
            error_response = {
                "jsonrpc": "2.0",
                "id": None,
                "error": {
                    "code": -32603,
                    "message": f"Internal error: {str(e)}"
                }
            }
            print(json.dumps(error_response))
            sys.stdout.flush()


if __name__ == "__main__":
    asyncio.run(main())
